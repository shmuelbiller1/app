"""File parsers: convert raw bytes of txt/csv/json/pdf/docx into a flat list of
text fragments (the atomic 'pieces of data' the optimizer dedups)."""
from __future__ import annotations

import csv
import io
import json
import re
from typing import List

_SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
_LINE_PREFIX = re.compile(r"^\s*(?:[-*•·▪◦]|\d+[.)]|[a-zA-Z][.)])\s+")


def _clean(line: str) -> str:
    line = line.replace("\u00a0", " ")
    line = _LINE_PREFIX.sub("", line.strip())
    return re.sub(r"\s+", " ", line).strip()


def split_text(raw: str, min_length: int = 3) -> List[str]:
    """Split free text into clean fragments (lines, then long lines by sentence)."""
    out: List[str] = []
    for rawline in raw.splitlines():
        line = _clean(rawline)
        if not line:
            continue
        if len(line) > 240 and _SENTENCE_SPLIT.search(line):
            for sent in _SENTENCE_SPLIT.split(line):
                sent = sent.strip()
                if len(sent) >= min_length:
                    out.append(sent)
        elif len(line) >= min_length:
            out.append(line)
    return out


def _flatten_json(obj, prefix: str = "") -> List[str]:
    rows: List[str] = []
    if isinstance(obj, dict):
        parts = []
        nested = []
        for k, v in obj.items():
            if isinstance(v, (dict, list)):
                nested.extend(_flatten_json(v, k))
            else:
                parts.append(f"{k}: {v}")
        if parts:
            rows.append(" | ".join(parts))
        rows.extend(nested)
    elif isinstance(obj, list):
        for item in obj:
            rows.extend(_flatten_json(item, prefix))
    else:
        rows.append(str(obj))
    return rows


def parse_bytes(data: bytes, filename: str, min_length: int = 3) -> List[str]:
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        return _parse_pdf(data, min_length)
    if name.endswith(".docx"):
        return _parse_docx(data, min_length)

    text = data.decode("utf-8", errors="replace")

    if name.endswith(".json"):
        try:
            obj = json.loads(text)
            rows = [r for r in _flatten_json(obj) if len(r) >= min_length]
            return rows or split_text(text, min_length)
        except Exception:
            return split_text(text, min_length)

    if name.endswith(".csv") or name.endswith(".tsv"):
        delim = "\t" if name.endswith(".tsv") else ","
        rows: List[str] = []
        reader = csv.reader(io.StringIO(text), delimiter=delim)
        header = None
        for i, cols in enumerate(reader):
            cols = [c.strip() for c in cols]
            if i == 0 and any(not _looks_numeric(c) for c in cols):
                header = cols
                continue
            if header and len(header) == len(cols):
                row = " | ".join(f"{h}: {c}" for h, c in zip(header, cols) if c)
            else:
                row = " | ".join(c for c in cols if c)
            if len(row) >= min_length:
                rows.append(row)
        return rows

    return split_text(text, min_length)


def _looks_numeric(s: str) -> bool:
    try:
        float(s)
        return True
    except ValueError:
        return False


def _parse_pdf(data: bytes, min_length: int) -> List[str]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    chunks: List[str] = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        chunks.extend(split_text(txt, min_length))
    return chunks


def _parse_docx(data: bytes, min_length: int) -> List[str]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    out: List[str] = []
    for para in doc.paragraphs:
        line = _clean(para.text)
        if len(line) >= min_length:
            out.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            joined = " | ".join(cells)
            if len(joined) >= min_length:
                out.append(joined)
    return out
